"""
提供.docx格式Word文件的处理功能
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from zipfile import BadZipFile
import os
import sys
import re

# 导入样式模块
from document_styles import (
    create_document_with_styles, 
    apply_title_style, 
    apply_subtitle_style, 
    apply_author_style, 
    apply_body_style,
    create_chinese_formal_document,
    apply_chinese_main_title,
    apply_chinese_subtitle,
    apply_chinese_body
)

# 导入工具模块
from file_utils import extract_author_from_filename, sanitize_filename, generate_output_filename
from image_extractor import extract_images_from_doc, extract_document_image_relations, find_paragraph_images

def process_docx_file(input_file, output_dir, suffix_enabled=True, 
                     suffix_text="——福州大学先进制造学院与海洋学院关工委2023年'中华魂'（毛泽东伟大精神品格）主题教育征文", 
                     use_chinese_format=False, keep_image_position=True, show_author_info=True,
                     mark_low_wordcount=False): # Added mark_low_wordcount parameter
    """
    处理单个.docx格式的Word文件
    
    参数:
        input_file: 输入文件路径
        output_dir: 输出目录
        suffix_enabled: 是否启用标题后缀
        suffix_text: 标题后缀内容
        use_chinese_format: 是否使用中文格式
        keep_image_position: 是否保持图片位置
        show_author_info: 是否显示作者信息
        mark_low_wordcount: 是否标记低字数文档
    
    返回:
        bool: 处理成功返回True，否则返回False
    """
    print(f"DEBUG: 开始处理文件 {input_file}")
    try:
        # 检查文件是否存在
        if not os.path.exists(input_file):
            print(f"× 错误：输入文件 '{input_file}' 不存在")
            return False

        # 创建成功文件文件夹
        success_dir = os.path.join(output_dir, "成功文件")
        no_image_dir = os.path.join(output_dir, "无图片成功文件")
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(no_image_dir, exist_ok=True)

        # 创建临时目录用于图片处理
        temp_dir = os.path.join(output_dir, "temp_images")
        os.makedirs(temp_dir, exist_ok=True)

        try:
            # 打开源文档
            source_doc = Document(input_file)
            
            # 创建目标文档
            if use_chinese_format:
                new_doc = create_chinese_formal_document()
            else:
                new_doc = create_document_with_styles()
                
            # 提取图片信息
            image_relations = extract_document_image_relations(source_doc)

            # 检查是否有图片
            has_images = bool(image_relations)
            print(f"DEBUG: 文档中包含 {len(image_relations)} 张图片")

            # 提取信息变量
            author_name = ""
            original_title = ""
            title_found = False
            author_added = False
            used_default_author = False
            
            # 从文件名中提取作者名
            filename = os.path.basename(input_file)
            author_name = extract_author_from_filename(filename)
            
            # 保存每个段落对应的图片，用于后续处理
            paragraph_images = {}
            
            # 首先遍历文档，找出所有带图片的段落
            if keep_image_position and has_images:
                print("DEBUG: 分析文档结构以保持图片位置...")
                paragraph_images = find_paragraph_images(source_doc, image_relations)
            
            # 处理文档内容
            for para_idx, para in enumerate(source_doc.paragraphs):
                try:
                    text = para.text.strip()
                    
                    # 对于空段落，检查是否有图片
                    if not text and para_idx in paragraph_images:
                        # 创建新段落并插入图片
                        img_para = new_doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        for img_path in paragraph_images[para_idx]:
                            try:
                                run = img_para.add_run()
                                run.add_picture(img_path, width=Inches(6))
                                print(f"DEBUG: 在空段落 {para_idx} 中添加图片")
                            except Exception as e:
                                print(f"× 添加图片时出错: {str(e)}")
                        
                        continue  # 处理完图片，继续下一段落
                    
                    if not text:
                        continue  # 跳过空段落
                    
                    # 提取标题（第一个非空段落）
                    if not title_found:
                        original_title = text
                        title_para = new_doc.add_paragraph(original_title)
                        if use_chinese_format:
                            apply_chinese_main_title(new_doc, title_para)
                        else:
                            apply_title_style(new_doc, title_para)
                        
                        # 添加副标题
                        if suffix_enabled and suffix_text:
                            subtitle_para = new_doc.add_paragraph(suffix_text)
                            if use_chinese_format:
                                apply_chinese_subtitle(new_doc, subtitle_para)
                            else:
                                apply_subtitle_style(new_doc, subtitle_para)
                        
                        title_found = True
                        continue
                    
                    # 从文档内容中提取作者信息
                    if not author_name and text.startswith('852'):
                        author_match = re.search(r'852\d*[^-]*-([^-\d\W]+)', text)
                        if not author_match:
                            author_match = re.search(r'852\d*[\s-]*([^\d\W]+)', text)
                        if author_match:
                            author_name = author_match.group(1).strip()
                        continue  # 跳过学号行
                    
                    # 添加作者信息（仅在需要且未添加时）
                    if author_name and not author_added and not text.startswith('852') and show_author_info:
                        author_text = f"（先进制造学院与海洋学院关工委通讯员{author_name}）"
                        author_para = new_doc.add_paragraph(author_text)
                        if use_chinese_format:
                            apply_chinese_subtitle(new_doc, author_para)
                        else:
                            apply_author_style(new_doc, author_para)
                        author_added = True
                        continue
                    
                    # 处理正文段落
                    if title_found and not text.startswith('852'):
                        # 添加文本段落
                        body_para = new_doc.add_paragraph(text)
                        if use_chinese_format:
                            apply_chinese_body(new_doc, body_para)
                        else:
                            apply_body_style(new_doc, body_para)
                        
                        # 如果段落有关联图片且需要保持图片位置，则添加图片
                        if keep_image_position and para_idx in paragraph_images:
                            # 添加段落后面的图片
                            img_para = new_doc.add_paragraph()
                            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            for img_path in paragraph_images[para_idx]:
                                try:
                                    run = img_para.add_run()
                                    run.add_picture(img_path, width=Inches(6))
                                    print(f"DEBUG: 在段落 {para_idx} 后添加图片")
                                except Exception as e:
                                    print(f"× 添加图片时出错: {str(e)}")
                except Exception as e:
                    print(f"× 处理段落 {para_idx} 时出错：{str(e)}")
                    continue

            # 处理默认作者名
            if not author_name:
                author_name = "佚名"
                used_default_author = True
                print(f"! 警告：未能提取作者名，使用默认值\"{author_name}\"")
                
                # 添加默认作者信息（如果需要且未添加）
                if not author_added and show_author_info and title_found:
                    author_text = f"（先进制造学院与海洋学院关工委通讯员{author_name}）"
                    author_para = new_doc.add_paragraph(author_text)
                    if use_chinese_format:
                        apply_chinese_subtitle(new_doc, author_para)
                    else:
                        apply_author_style(new_doc, author_para)
                    author_added = True

            # 如果不保持图片位置或没找到图片位置信息，但文档包含图片，则在末尾添加图片
            if has_images and not keep_image_position:
                print("DEBUG: 将所有图片添加到文档末尾")
                new_doc.add_paragraph()  # 添加空行分隔
                
                # 将所有图片提取并添加到末尾
                temp_img_dir = os.path.join(temp_dir, "all_images")
                os.makedirs(temp_img_dir, exist_ok=True)
                
                image_files = extract_images_from_doc(input_file, temp_img_dir)
                
                for img_path in image_files:
                    try:
                        img_para = new_doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = img_para.add_run()
                        run.add_picture(img_path, width=Inches(6))
                    except Exception as e:
                        print(f"× 添加图片到文档末尾时出错：{str(e)}")

            # 生成输出文件名
            if original_title:
                # 生成文件名
                new_filename = generate_output_filename(
                    author_name, original_title, suffix_enabled, suffix_text, show_author_info,
                    mark_low_wordcount=mark_low_wordcount # Pass mark_low_wordcount
                )
                
                # 根据是否有图片选择保存目录
                output_dir_final = success_dir if has_images else no_image_dir
                output_file = os.path.join(output_dir_final, new_filename)
                
                # 保存文档
                try:
                    new_doc.save(output_file)
                    if used_default_author:
                        if has_images:
                            print(f"✓ 文件处理完成（使用默认作者名）：{new_filename}")
                        else:
                            print(f"✓ 文件处理完成（使用默认作者名，无图片）：{new_filename}")
                    else:
                        if has_images:
                            print(f"✓ 文件处理完成：{new_filename}")
                        else:
                            print(f"✓ 文件处理完成（无图片）：{new_filename}")
                except Exception as e:
                    print(f"× 保存文件时出错：{str(e)}")
                    return False
            else:
                print("× 未能提取标题")
                return False
                
        except BadZipFile:
            print(f"× 错误：文件 '{input_file}' 可能已损坏或不是有效的Word文档")
            return False

        # 清理临时文件
        try:
            if os.path.exists(temp_dir):
                import shutil
                shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"! 清理临时文件时出错: {str(e)}")

        return True

    except Exception as e:
        print(f"× 处理文件时出现错误：{str(e)}")
        return False