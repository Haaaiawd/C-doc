"""
提供标题提取相关的业务逻辑处理
"""
import os
import threading
from docx import Document
from datetime import datetime
from word_processors import extract_author_from_filename

# 检查是否安装了openpyxl库
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class TitleHandler:
    """标题提取处理器，处理从文档中提取标题的业务逻辑"""
    
    def __init__(self, app):
        """
        初始化标题提取处理器
        
        参数:
            app: 主应用程序实例
        """
        self.app = app
        
    def validate_paths(self, input_dir, output_dir=None):
        """
        验证路径是否有效
        
        参数:
            input_dir: 输入目录
            output_dir: 输出目录(可选)
            
        返回:
            tuple: (是否有效, 错误信息)
        """
        if not input_dir:
            return False, "请选择输入文件夹"
            
        # 检查输入目录是否存在
        if not os.path.exists(input_dir):
            return False, "输入目录不存在"
        
        # 如果提供了输出目录，则检查
        if output_dir is not None and not output_dir:
            return False, "请选择输出文件夹"
            
        return True, ""
        
    def extract_titles(self):
        """开始提取文档标题"""
        # 获取路径配置
        paths = self.app.file_frame.get_paths()
        input_dir = paths["input_dir"]
        output_dir = paths["output_dir"]
        
        # 验证路径
        valid, error_msg = self.validate_paths(input_dir, output_dir)
        if not valid:
            self.app.set_status(error_msg)
            return
        
        # 准备开始提取标题
        self.app.reset_for_processing()
        self.app.set_status("正在提取标题...")
        
        # 在新线程中运行标题提取
        thread = threading.Thread(
            target=self._extract_thread, 
            args=(input_dir, output_dir),
            daemon=True
        )
        thread.start()
    
    def _extract_thread(self, input_dir, output_dir):
        """标题提取线程"""
        try:
            # 获取所有Word文件
            docx_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]
            total_files = len(docx_files)
            title_data = []
            
            # 配置进度条最大值
            self.app.root.after(0, lambda: self.app.progress_bar.configure(maximum=total_files))
            
            for index, filename in enumerate(docx_files, 1):
                # 更新进度条
                current_progress = index / total_files * 100
                self.app.root.after(0, lambda p=index: self.app.progress_bar.configure(value=p))
                self.app.root.after(0, lambda: self.app.progress_bar.update())
                
                self.app.set_status(f"正在提取标题: {index}/{total_files} ({int(current_progress)}%)...")
                
                input_file = os.path.join(input_dir, filename)
                try:
                    doc = Document(input_file)
                    
                    # 尝试从文档属性获取标题
                    try:
                        core_props = doc.core_properties
                        doc_title = core_props.title or ""
                    except:
                        doc_title = ""
                    
                    # 如果文档属性中没有标题，从文档内容中提取
                    if not doc_title and len(doc.paragraphs) > 0:
                        # 检查第一段是否为标题（通常是加粗或大字体）
                        first_para = doc.paragraphs[0]
                        if len(first_para.runs) > 0 and first_para.text.strip():
                            if first_para.runs[0].bold or first_para.runs[0].font.size is not None:
                                doc_title = first_para.text.strip()
                    
                    # 如果仍未找到标题，使用第一个非空段落
                    if not doc_title:
                        for para in doc.paragraphs:
                            if para.text.strip():
                                doc_title = para.text.strip()
                                break
                    
                    # 如果标题太长，可能是摘要或正文，截取合理长度
                    if len(doc_title) > 100:
                        doc_title = doc_title[:97] + "..."
                    
                    # 获取作者信息
                    author = extract_author_from_filename(filename) or "未知"
                    
                    title_data.append((filename, doc_title, author))
                    self.app.log_text.insert('end', f"{filename} ({author}): {doc_title}\n")
                    
                except Exception as e:
                    self.app.log_text.insert('end', f"× {filename}: 提取标题失败 - {str(e)}\n")
            
            # 导出为Excel
            if title_data and EXCEL_AVAILABLE:
                try:
                    excel_path = self._export_titles_excel(title_data, output_dir)
                    self.app.log_text.insert('end', f"\nExcel报告已导出至: {excel_path}\n")
                except Exception as e:
                    self.app.log_text.insert('end', f"\nExcel报告导出失败: {str(e)}\n")
            
            self.app.log_text.see('end')
            
            # 设置进度条为100%完成
            self.app.root.after(0, lambda: self.app.progress_bar.configure(value=total_files))
            
            # 更新状态
            self.app.root.after(0, lambda: self.app.set_status(f"标题提取完成，共 {len(title_data)} 个文件"))
            self.app.root.after(0, self.app.enable_buttons)
            
        except Exception as e:
            self.app.root.after(0, lambda: self.app.set_status(f"提取标题时出错: {str(e)}"))
            self.app.root.after(0, self.app.enable_buttons)
    
    def _export_titles_excel(self, title_data, output_dir):
        """导出标题列表到Excel"""
        if not EXCEL_AVAILABLE:
            raise ImportError("缺少openpyxl库，无法导出Excel")
        
        # 创建报告目录
        reports_dir = os.path.join(output_dir, "标题提取报告")
        os.makedirs(reports_dir, exist_ok=True)
        
        # 创建带有当前日期时间的文件名
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(reports_dir, f"标题提取报告_{current_time}.xlsx")
        
        # 创建新的Excel工作簿
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = "文档标题列表"
        
        # 设置标题行
        headers = ["序号", "文件名", "标题", "作者"]
        for col, header in enumerate(headers, 1):
            cell = sheet.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
        
        # 写入数据
        for row, (filename, title, author) in enumerate(title_data, 2):
            sheet.cell(row=row, column=1).value = row - 1  # 序号
            sheet.cell(row=row, column=2).value = filename
            sheet.cell(row=row, column=3).value = title
            sheet.cell(row=row, column=4).value = author
        
        # 自动调整列宽
        try:
            for sheet in wb.worksheets:
                for col in range(1, sheet.max_column + 1):
                    max_length = 0
                    column = get_column_letter(col)
                    
                    # 计算最大列宽
                    for cell in sheet[column]:
                        try:
                            if cell.value:
                                cell_length = len(str(cell.value))
                                if cell_length > max_length:
                                    max_length = cell_length
                        except:
                            pass
                    
                    # 设置列宽，考虑中文字符宽度
                    adjusted_width = (max_length + 2) * 1.2
                    sheet.column_dimensions[column].width = adjusted_width
        except Exception as e:
            print(f"! 调整Excel列宽时出错: {str(e)}")
        
        # 保存工作簿
        try:
            wb.save(excel_path)
        except Exception as e:
            # 尝试使用备用文件名
            alt_path = os.path.join(
                os.path.dirname(excel_path),
                f"标题提取报告_{current_time}_新.xlsx"
            )
            wb.save(alt_path)
            excel_path = alt_path
        
        return excel_path