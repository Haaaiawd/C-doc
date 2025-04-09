import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
import threading
from process_word import (
    process_word_file, 
    extract_author_number, 
    extract_author_from_filename,
    extract_images_from_doc,
    process_doc_file  # 添加导入process_doc_file函数
)
from heading_utils import count_document_words, get_document_stats
import sys
import io
import os
import shutil
import re
from docx import Document
# 新增依赖库
from datetime import datetime
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

class RedirectText:
    def __init__(self, text_widget, error_only=False):
        self.text_widget = text_widget
        self.error_only = error_only
        self.error_files = set()  # 存储错误文件路径
        self.success_files = set()  # 存储成功文件路径
        self.current_file = None  # 当前正在处理的文件

    def set_current_file(self, filename):
        self.current_file = filename

    def write(self, string):
        if not string.strip():
            return
        # 检查是否是错误信息
        if string.strip().startswith('×'):
            # 如果有当前文件，添加到错误文件集合中
            if self.current_file:
                self.error_files.add(self.current_file)
        # 检查是否是成功信息（包括使用默认作者名的成功处理）
        elif string.strip().startswith('✓'):
            # 如果有当前文件，添加到成功文件集合中
            if self.current_file:
                self.success_files.add(self.current_file)
            
        # 显示错误信息，每个文件占一行
        if string.strip().startswith('×'):
            if self.current_file:
                if not string.strip().endswith('.docx'):
                    # 提取作者行数字和作者名
                    author_num = extract_author_number(self.current_file)
                    author_name = extract_author_from_filename(self.current_file)
                    error_msg = f"作者{author_num}({author_name}): {string.strip()}\n"
                else:
                    error_msg = string
            else:
                error_msg = string
            
            self.text_widget.insert('end', error_msg)
            self.text_widget.see('end')
            self.text_widget.update()
        # 显示警告信息，每个文件占一行
        elif string.strip().startswith('!'):
            if self.current_file:
                if not string.strip().endswith('.docx'):
                    # 提取作者行数字
                    author_num = extract_author_number(self.current_file)
                    warning_msg = f"作者{author_num}: {string.strip()}\n"
                else:
                    warning_msg = string
            else:
                warning_msg = string
            
            self.text_widget.insert('end', warning_msg)
            self.text_widget.see('end')
            self.text_widget.update()

    def get_error_files(self):
        return self.error_files

    def get_success_files(self):
        return self.success_files

    def clear_files(self):
        self.error_files.clear()
        self.success_files.clear()
        self.current_file = None

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Word档批量处理工具")
        self.root.geometry("800x650")  # 增加窗口高度以容纳新控件
        
        # 创建主框架
        main_frame = ttk.Frame(root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 输入文件夹选择
        ttk.Label(main_frame, text="输入文件夹:").grid(row=0, column=0, sticky=tk.W)
        self.input_path = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.input_path, width=50).grid(row=0, column=1, padx=5)
        ttk.Button(main_frame, text="浏览", command=self.choose_input_dir).grid(row=0, column=2)
        
        # 输出文件夹选择
        ttk.Label(main_frame, text="输出文件夹:").grid(row=1, column=0, sticky=tk.W)
        self.output_path = tk.StringVar()
        ttk.Entry(main_frame, textvariable=self.output_path, width=50).grid(row=1, column=1, padx=5)
        ttk.Button(main_frame, text="浏览", command=self.choose_output_dir).grid(row=1, column=2)
        
        # 标题后缀选项
        ttk.Label(main_frame, text="标题后缀:").grid(row=2, column=0, sticky=tk.W)
        self.suffix_enabled = tk.BooleanVar(value=True)
        ttk.Checkbutton(main_frame, text="启用标题后缀", variable=self.suffix_enabled).grid(row=2, column=2)
        
        # 标题后缀文本框
        self.suffix_text = tk.StringVar(value="——福州大学先进制造学院与海洋学院关工委2023年'中华魂'（毛泽东伟大精神品格）主题教育征文")
        ttk.Entry(main_frame, textvariable=self.suffix_text, width=50).grid(row=2, column=1, padx=5)
        
        # 添加格式选择组
        format_frame = ttk.LabelFrame(main_frame, text="文档格式")
        format_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), padx=5, pady=5)
        
        # 创建格式单选按钮
        self.format_var = tk.StringVar(value="default")
        ttk.Radiobutton(format_frame, text="默认格式", variable=self.format_var, value="default").grid(row=0, column=0, padx=10)
        ttk.Radiobutton(format_frame, text="中文标准格式", variable=self.format_var, value="chinese").grid(row=0, column=1, padx=10)
        
        # 中文标准格式说明
        format_desc = tk.Text(format_frame, height=4, width=70, wrap=tk.WORD)
        format_desc.grid(row=1, column=0, columnspan=2, padx=5, pady=5)
        format_desc.insert(tk.END, "中文标准格式详情：\n主标题：小二号方正小标宋简体\n副标题：小三号仿宋_GB2312\n正文：小三号仿宋_GB2312\n文内一级标题：黑体小三，二级标题：楷体_GB2312小三，三四级标题：仿宋_GB2312小三")
        format_desc.config(state=tk.DISABLED)  # 设为只读
        
        # 添加保持图片位置选项
        self.keep_image_position = tk.BooleanVar(value=True)
        ttk.Checkbutton(format_frame, text="保持图片在原位置", variable=self.keep_image_position).grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        
        # 添加作者信息控制选项
        self.show_author_info = tk.BooleanVar(value=True)
        ttk.Checkbutton(format_frame, text="添加作者信息（先进制造学院与海洋学院关工委通讯员）", 
                         variable=self.show_author_info).grid(row=2, column=1, sticky=tk.W, padx=10, pady=5)
        
        # 添加字数检测框架
        wordcount_frame = ttk.LabelFrame(main_frame, text="字数检测")
        wordcount_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), padx=5, pady=5)
        
        # 添加字数检测启用选项
        self.wordcount_enabled = tk.BooleanVar(value=False)
        ttk.Checkbutton(wordcount_frame, text="启用字数检测", variable=self.wordcount_enabled).grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        
        # 添加最小字数设置
        ttk.Label(wordcount_frame, text="最小字数要求:").grid(row=0, column=1, sticky=tk.E, padx=5, pady=2)
        self.min_words = tk.IntVar(value=800)  # 默认最小字数为800
        ttk.Spinbox(wordcount_frame, from_=100, to=5000, increment=100, textvariable=self.min_words, width=5).grid(row=0, column=2, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(wordcount_frame, text="字").grid(row=0, column=3, sticky=tk.W, padx=0, pady=2)
        
        # 添加字数不足文件处理选项
        ttk.Label(wordcount_frame, text="字数不足的文件:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        self.wordcount_action = tk.StringVar(value="mark")
        ttk.Radiobutton(wordcount_frame, text="标记但仍处理", variable=self.wordcount_action, value="mark").grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        ttk.Radiobutton(wordcount_frame, text="移至单独文件夹", variable=self.wordcount_action, value="move").grid(row=1, column=2, columnspan=2, sticky=tk.W, padx=5, pady=2)
        
        # 添加导出Excel选项
        self.export_excel = tk.BooleanVar(value=True)
        excel_check = ttk.Checkbutton(wordcount_frame, text="导出Excel报告", variable=self.export_excel)
        excel_check.grid(row=2, column=0, sticky=tk.W, padx=5, pady=2)
        if not EXCEL_AVAILABLE:
            excel_check.state(['disabled'])
            self.export_excel.set(False)
            ttk.Label(wordcount_frame, text="(需要安装openpyxl库)", foreground="red").grid(row=2, column=1, sticky=tk.W, padx=5, pady=2)
        
        # 按钮框架
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=5, column=0, columnspan=3, pady=10)  # 修改行号
        
        # 转换按钮
        self.convert_btn = ttk.Button(button_frame, text="开始转换", command=self.start_conversion)
        self.convert_btn.pack(side=tk.LEFT, padx=5)
        
        # 打包错误文件按钮
        self.pack_error_btn = ttk.Button(button_frame, text="打包错误文件", command=self.pack_error_files)
        self.pack_error_btn.pack(side=tk.LEFT, padx=5)
        self.pack_error_btn.state(['disabled'])
        
        # 添加提取图片按钮
        self.extract_images_btn = ttk.Button(button_frame, text="提取图片", command=self.extract_images)
        self.extract_images_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加提取标题按钮
        self.extract_titles_btn = ttk.Button(button_frame, text="提取标题", command=self.extract_titles)
        self.extract_titles_btn.pack(side=tk.LEFT, padx=5)
        
        # 添加检测字数按钮
        self.check_wordcount_btn = ttk.Button(button_frame, text="检测字数", command=self.check_wordcount)
        self.check_wordcount_btn.pack(side=tk.LEFT, padx=5)
        
        # 进度显示
        self.progress_var = tk.StringVar(value="就绪")
        ttk.Label(main_frame, textvariable=self.progress_var).grid(row=6, column=0, columnspan=3)  # 修改行号
        
        # 错误信息显示区域
        ttk.Label(main_frame, text="信息显示:").grid(row=7, column=0, sticky=tk.W)  # 修改行号
        self.log_text = scrolledtext.ScrolledText(main_frame, height=20, width=80)
        self.log_text.grid(row=8, column=0, columnspan=3, pady=5)  # 修改行号
        
        # 配置grid权重
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # 重定向标准输出到文本框
        self.redirect = RedirectText(self.log_text, error_only=True)
        sys.stdout = self.redirect
        sys.stderr = self.redirect

    def choose_input_dir(self):
        directory = filedialog.askdirectory()
        if directory:
            self.input_path.set(directory)
            # 默认设置相同的输出目录
            if not self.output_path.get():
                self.output_path.set(directory)

    def choose_output_dir(self):
        directory = filedialog.askdirectory()
        if directory:
            self.output_path.set(directory)

    def pack_error_files(self):
        error_files = self.redirect.get_error_files()
        if not error_files:
            self.progress_var.set("没有错误文件需要打包")
            return
        
        try:
            # 创建错误文件文件夹
            error_dir = os.path.join(self.output_path.get(), "错误文件")
            os.makedirs(error_dir, exist_ok=True)
            
            # 检查磁盘空间
            if shutil.disk_usage(error_dir).free < 1024 * 1024 * 100:  # 100MB
                self.progress_var.set("磁盘空间不足")
                return
            
            # 复制错误文件
            copied_count = 0
            for filename in error_files:
                try:
                    src_path = os.path.join(self.input_path.get(), filename)
                    if os.path.exists(src_path):
                        dst_path = os.path.join(error_dir, filename)
                        shutil.copy2(src_path, dst_path)
                        copied_count += 1
                except Exception as e:
                    print(f"× 复制文件 {filename} 时出错：{str(e)}")
                
            self.progress_var.set(f"已将 {copied_count} 个错误文件复制到 {error_dir}")
            # 复制完成后禁用打包按钮
            self.pack_error_btn.state(['disabled'])
            
        except PermissionError:
            self.progress_var.set("没有足够的权限创建或访问目录")
        except Exception as e:
            self.progress_var.set(f"打包错误文件时出错: {str(e)}")

    def start_conversion(self):
        input_dir = self.input_path.get()
        output_dir = self.output_path.get()
        
        if not input_dir or not output_dir:
            self.progress_var.set("请选择输入和输出文件夹")
            return
        
        # 检查目录是否存在
        if not os.path.exists(input_dir):
            self.progress_var.set("输入目录不存在")
            return
        
        # 检查是否有 Word 文件
        input_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]  # 修改这里添加.doc扩展名
        if not input_files:
            self.progress_var.set("输入目录中没有 Word 文件")
            return
        
        # 清空之前的文件记录
        self.redirect.clear_files()
        
        # 禁用所有按钮
        self.convert_btn.state(['disabled'])
        self.pack_error_btn.state(['disabled'])
        self.check_wordcount_btn.state(['disabled'])
        self.progress_var.set("处理中...")
        self.log_text.delete(1.0, tk.END)
        
        # 获取标题后缀配置
        suffix_enabled = self.suffix_enabled.get()
        suffix_text = self.suffix_text.get() if suffix_enabled else ""
        
        # 获取格式选择
        use_chinese_format = self.format_var.get() == "chinese"
        
        # 获取图片位置和作者信息选项
        keep_image_position = self.keep_image_position.get()
        show_author_info = self.show_author_info.get()
        
        # 获取字数检测配置
        wordcount_enabled = self.wordcount_enabled.get()
        min_words = self.min_words.get()
        wordcount_action = self.wordcount_action.get()
        
        # 向用户显示所选配置信息
        format_info = "中文标准格式" if use_chinese_format else "默认格式"
        self.log_text.insert('end', f"使用{format_info}处理文档\n")
        
        if not show_author_info:
            self.log_text.insert('end', "不添加作者信息\n")
            
        if keep_image_position:
            self.log_text.insert('end', "保持图片在原文位置\n")
        else:
            self.log_text.insert('end', "所有图片将移至文末\n")
        
        self.log_text.insert('end', "\n开始处理文件...\n\n")
        
        # 创建临时目录
        temp_dir = os.path.join(output_dir, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        # 创建字数不足文件夹
        if wordcount_enabled and wordcount_action == "move":
            low_wordcount_dir = os.path.join(output_dir, f"字数不足{min_words}字")
            os.makedirs(low_wordcount_dir, exist_ok=True)
        
        # 在新线程中运行转换
        def conversion_thread():
            try:
                # 获取目录中的所有文件并排序
                input_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]  # 修改这里添加.doc扩展名
                sorted_files = sorted(input_files, key=lambda x: extract_author_number(x))
                
                total_files = len(sorted_files)
                self.progress_var.set(f"开始处理，共 {total_files} 个文件...")
                
                # 字数统计结果
                low_wordcount_files = []
                
                # 处理每个文件
                for index, filename in enumerate(sorted_files, 1):
                    self.progress_var.set(f"正在处理第 {index}/{total_files} 个文件...")
                    input_file = os.path.join(input_dir, filename)
                    
                    # 设置当前处理的文件
                    self.redirect.set_current_file(filename)
                    
                    # 如果启用了字数检测，先检查字数
                    if wordcount_enabled:
                        try:
                            doc = Document(input_file)
                            word_count = count_document_words(doc)
                            
                            # 字数不足
                            if word_count < min_words:
                                self.log_text.insert('end', f"! {filename}: 字数为 {word_count}，不足 {min_words} 字\n")
                                low_wordcount_files.append((filename, word_count))
                                
                                # 如果选择移动字数不足文件到单独文件夹
                                if wordcount_action == "move":
                                    target_file = os.path.join(low_wordcount_dir, filename)
                                    shutil.copy2(input_file, target_file)
                                    continue  # 跳过后续处理
                            else:
                                self.log_text.insert('end', f"✓ {filename}: 字数为 {word_count}，满足要求\n")
                        except Exception as e:
                            self.log_text.insert('end', f"! {filename}: 字数检测失败 - {str(e)}\n")
                    
                    # 根据文件扩展名选择处理方法，传入标题后缀配置、格式选择以及新增的选项
                    if filename.lower().endswith('.doc'):
                        process_doc_file(input_file, output_dir, suffix_enabled, suffix_text, use_chinese_format, 
                                         keep_image_position, show_author_info)
                    else:  # .docx 文件
                        process_word_file(input_file, output_dir, suffix_enabled, suffix_text, use_chinese_format,
                                         keep_image_position, show_author_info)
                
                # 处理完成后显示统计信息
                success_count = len(self.redirect.get_success_files())
                failed_count = len(self.redirect.get_error_files())
                low_count = len(low_wordcount_files)
                
                summary = f"\n处理完成！\n总计: {total_files} 个文件\n成功: {success_count} 个\n失败: {failed_count} 个\n"
                
                if wordcount_enabled:
                    summary += f"字数不足: {low_count} 个\n"
                    if low_count > 0:
                        summary += "\n字数不足的文件:\n"
                        for filename, count in low_wordcount_files:
                            summary += f"{filename}: {count} 字\n"
                
                if failed_count > 0:
                    summary += "\n失败的文件作者数字:\n"
                    for failed_file in self.redirect.get_error_files():
                        author_num = extract_author_number(failed_file)
                        summary += f"作者{author_num}\n"
                
                self.log_text.insert('end', summary)
                self.log_text.see('end')
                
                self.root.after(0, self.conversion_complete)
            except Exception as e:
                self.root.after(0, lambda: self.conversion_error(str(e)))
            
            # 处理完成后清理临时目录
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                print(f"! 清理临时文件时出错: {str(e)}")
        
        threading.Thread(target=conversion_thread, daemon=True).start()

    def conversion_complete(self):
        success_count = len(self.redirect.get_success_files())
        failed_count = len(self.redirect.get_error_files())
        self.progress_var.set(f"处理完成，成功: {success_count} 个，失败: {failed_count} 个")
        self.convert_btn.state(['!disabled'])
        self.check_wordcount_btn.state(['!disabled'])
        if failed_count > 0:
            self.pack_error_btn.state(['!disabled'])

    def conversion_error(self, error_message):
        self.progress_var.set(f"处理出错: {error_message}")
        self.convert_btn.state(['!disabled'])
        self.check_wordcount_btn.state(['!disabled'])
        # 只检查错误文件
        if self.redirect.get_error_files():
            self.pack_error_btn.state(['!disabled'])

    def extract_images(self):
        input_dir = self.input_path.get()
        output_dir = self.output_path.get()
        if not input_dir or not output_dir:
            self.progress_var.set("请选择输入和输出文件夹")
            return
        
        if not os.path.exists(input_dir):
            self.progress_var.set("输入目录不存在")
            return
        
        # 清空显示框
        self.log_text.delete(1.0, tk.END)
        self.progress_var.set("正在提取图片...")
        self.extract_images_btn.state(['disabled'])
        
        def extract_thread():
            try:
                # 创建图片输出目录
                images_dir = os.path.join(self.output_path.get(), "提取的图片")
                os.makedirs(images_dir, exist_ok=True)
                
                # 获取所有Word文件
                docx_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]  # 修改这里添加.doc扩展名
                total_images = 0
                
                for filename in docx_files:
                    input_file = os.path.join(input_dir, filename)
                    # 为每个文件创建子文件夹
                    file_images_dir = os.path.join(images_dir, os.path.splitext(filename)[0])
                    os.makedirs(file_images_dir, exist_ok=True)
                    
                    # 提取图片
                    temp_images = extract_images_from_doc(input_file, file_images_dir)
                    if temp_images:
                        self.log_text.insert('end', f"✓ {filename}: 提取了 {len(temp_images)} 张图片\n")
                        total_images += len(temp_images)
                    else:
                        self.log_text.insert('end', f"! {filename}: 未找到图片\n")
                    
                summary = f"\n提取完成！\n总计处理 {len(docx_files)} 个文件\n共提取 {total_images} 张图片\n"
                self.log_text.insert('end', summary)
                self.log_text.see('end')
                
                self.root.after(0, lambda: self.progress_var.set(f"图片提取完成，共 {total_images} 张"))
                self.root.after(0, lambda: self.extract_images_btn.state(['!disabled']))
                
            except Exception as e:
                self.root.after(0, lambda: self.progress_var.set(f"提取图片时出错: {str(e)}"))
                self.root.after(0, lambda: self.extract_images_btn.state(['!disabled']))
        
        threading.Thread(target=extract_thread, daemon=True).start()

    def extract_titles(self):
        input_dir = self.input_path.get()
        if not input_dir:
            self.progress_var.set("请选择输入文件夹")
            return
        
        if not os.path.exists(input_dir):
            self.progress_var.set("输入目录不存在")
            return
        
        # 清空显示框
        self.log_text.delete(1.0, tk.END)
        self.progress_var.set("正在提取标题...")
        self.extract_titles_btn.state(['disabled'])
        
        def extract_thread():
            try:
                # 获取所有Word文件
                docx_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]  # 修改这里添加.doc扩展名
                titles = []
                
                for filename in docx_files:
                    input_file = os.path.join(input_dir, filename)
                    try:
                        doc = Document(input_file)
                        # 获取第一个非空段落作为标题
                        title_found = False
                        for para in doc.paragraphs:
                            if para.text.strip():
                                titles.append((filename, para.text.strip()))
                                title_found = True
                                break
                        # 如果没有找到标题，使用默认值
                        if not title_found:
                            titles.append((filename, "无标题"))
                            self.log_text.insert('end', f"! {filename}: 未找到标题，使用默认值\"无标题\"\n")
                    except Exception as e:
                        self.log_text.insert('end', f"× {filename}: 提取标题失败 - {str(e)}\n")
                
                # 按文件名排序
                titles.sort(key=lambda x: extract_author_number(x[0]))
                
                # 显示标题
                self.log_text.insert('end', "提取的标题：\n" + "="*50 + "\n\n")
                for filename, title in titles:
                    self.log_text.insert('end', f"【{filename}】\n{title}\n\n")
                
                summary = f"\n提取完成！\n总计处理 {len(docx_files)} 个文件\n成功提取 {len(titles)} 个标题\n"
                self.log_text.insert('end', "="*50 + "\n" + summary)
                self.log_text.see('end')
                
                self.root.after(0, lambda: self.progress_var.set(f"标题提取完成，共 {len(titles)} 个"))
                self.root.after(0, lambda: self.extract_titles_btn.state(['!disabled']))
                
            except Exception as e:
                self.root.after(0, lambda: self.progress_var.set(f"提取标题时出错: {str(e)}"))
                self.root.after(0, lambda: self.extract_titles_btn.state(['!disabled']))
        
        threading.Thread(target=extract_thread, daemon=True).start()

    # 添加字数检测功能
    def check_wordcount(self):
        input_dir = self.input_path.get()
        output_dir = self.output_path.get()
        
        if not input_dir:
            self.progress_var.set("请选择输入文件夹")
            return
        
        if not output_dir:
            self.progress_var.set("请选择输出文件夹")
            return
        
        if not os.path.exists(input_dir):
            self.progress_var.set("输入目录不存在")
            return
        
        # 清空显示框
        self.log_text.delete(1.0, tk.END)
        self.progress_var.set("正在检测字数...")
        self.check_wordcount_btn.state(['disabled'])
        
        # 获取字数要求
        min_words = self.min_words.get()
        # 是否导出Excel
        export_excel = self.export_excel.get() and EXCEL_AVAILABLE
        
        def check_thread():
            try:
                # 获取所有Word文件
                docx_files = [f for f in os.listdir(input_dir) if f.endswith(('.doc', '.docx'))]
                word_counts = []
                low_wordcount_files = []
                
                total_files = len(docx_files)
                current = 0
                
                for filename in docx_files:
                    current += 1
                    self.progress_var.set(f"正在检测文件 {current}/{total_files}...")
                    input_file = os.path.join(input_dir, filename)
                    try:
                        doc = Document(input_file)
                        stats = get_document_stats(doc)
                        word_count = stats['word_count']
                        para_count = stats['paragraph_count']
                        char_count = stats['character_count']
                        
                        # 提取作者名
                        author_name = extract_author_from_filename(filename) or "未知"
                        
                        word_counts.append((filename, word_count, para_count, char_count, author_name))
                        if word_count < min_words:
                            low_wordcount_files.append((filename, word_count, author_name))
                    except Exception as e:
                        self.log_text.insert('end', f"× {filename}: 字数检测失败 - {str(e)}\n")
                
                # 按字数排序
                word_counts.sort(key=lambda x: x[1])
                
                # 显示结果
                self.log_text.insert('end', f"字数统计报告 (最小字数要求: {min_words})\n")
                self.log_text.insert('end', "="*50 + "\n\n")
                
                # 显示所有文件的字数
                self.log_text.insert('end', "所有文件的字数统计:\n")
                for filename, word_count, para_count, char_count, author_name in word_counts:
                    status = "不足" if word_count < min_words else "合格"
                    self.log_text.insert('end', f"{filename} ({author_name}): {word_count} 字 ({para_count} 段落, {char_count} 字符) - {status}\n")
                
                # 显示字数不足的文件
                if low_wordcount_files:
                    self.log_text.insert('end', "\n\n字数不足的文件:\n")
                    for filename, word_count, author_name in low_wordcount_files:
                        self.log_text.insert('end', f"{filename} ({author_name}): {word_count} 字\n")
                
                # 统计信息
                summary = f"\n\n统计信息:\n"
                summary += f"总文件数: {len(word_counts)} 个\n"
                summary += f"平均字数: {sum(x[1] for x in word_counts) / len(word_counts) if word_counts else 0:.1f} 字\n"
                summary += f"最大字数: {max(x[1] for x in word_counts) if word_counts else 0} 字\n"
                summary += f"最小字数: {min(x[1] for x in word_counts) if word_counts else 0} 字\n"
                summary += f"字数不足文件: {len(low_wordcount_files)} 个\n"
                summary += f"字数合格文件: {len(word_counts) - len(low_wordcount_files)} 个\n"
                
                self.log_text.insert('end', summary)
                
                # 导出Excel报告
                excel_path = ""
                if export_excel and word_counts:
                    try:
                        excel_path = self.export_wordcount_excel(word_counts, low_wordcount_files, min_words, output_dir)
                        self.log_text.insert('end', f"\n\nExcel报告已导出至: {excel_path}\n")
                    except Exception as e:
                        self.log_text.insert('end', f"\n\nExcel报告导出失败: {str(e)}\n")
                
                self.log_text.see('end')
                
                self.root.after(0, lambda: self.progress_var.set(
                    f"字数检测完成，共 {len(word_counts)} 个文件，其中 {len(low_wordcount_files)} 个不足 {min_words} 字" +
                    (f"，已生成Excel报告" if excel_path else "")
                ))
                self.root.after(0, lambda: self.check_wordcount_btn.state(['!disabled']))
                
            except Exception as e:
                self.root.after(0, lambda: self.progress_var.set(f"字数检测出错: {str(e)}"))
                self.root.after(0, lambda: self.check_wordcount_btn.state(['!disabled']))
        
        threading.Thread(target=check_thread, daemon=True).start()
    
    def export_wordcount_excel(self, word_counts, low_wordcount_files, min_words, output_dir):
        """导出字数统计Excel报告"""
        if not EXCEL_AVAILABLE:
            raise ImportError("缺少openpyxl库，无法导出Excel")
        
        # 创建报告目录
        reports_dir = os.path.join(output_dir, "字数统计报告")
        os.makedirs(reports_dir, exist_ok=True)
        
        # 创建带有当前日期时间的文件名
        current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        excel_path = os.path.join(reports_dir, f"字数统计报告_{current_time}.xlsx")
        
        # 创建新的Excel工作簿
        wb = openpyxl.Workbook()
        
        # 创建总览工作表
        overview_sheet = wb.active
        overview_sheet.title = "字数统计总览"
        
        # 设置标题行
        headers = ["文件名", "作者", "字数", "段落数", "字符数", "状态"]
        for col, header in enumerate(headers, 1):
            cell = overview_sheet.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
        
        # 写入数据
        for row, (filename, word_count, para_count, char_count, author_name) in enumerate(word_counts, 2):
            status = "不足" if word_count < min_words else "合格"
            
            overview_sheet.cell(row=row, column=1).value = filename
            overview_sheet.cell(row=row, column=2).value = author_name
            overview_sheet.cell(row=row, column=3).value = word_count
            overview_sheet.cell(row=row, column=4).value = para_count
            overview_sheet.cell(row=row, column=5).value = char_count
            overview_sheet.cell(row=row, column=6).value = status
            
            # 设置字数不足的行底色为浅红色
            if word_count < min_words:
                for col in range(1, 7):
                    overview_sheet.cell(row=row, column=col).fill = PatternFill(
                        start_color="FFCCCC", end_color="FFCCCC", fill_type="solid"
                    )
        
        # 创建字数不足工作表
        if low_wordcount_files:
            low_sheet = wb.create_sheet(title="字数不足文件")
            
            # 设置标题行
            headers = ["文件名", "作者", "字数", "差额"]
            for col, header in enumerate(headers, 1):
                cell = low_sheet.cell(row=1, column=col)
                cell.value = header
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
            
            # 写入数据
            for row, (filename, word_count, author_name) in enumerate(low_wordcount_files, 2):
                low_sheet.cell(row=row, column=1).value = filename
                low_sheet.cell(row=row, column=2).value = author_name
                low_sheet.cell(row=row, column=3).value = word_count
                low_sheet.cell(row=row, column=4).value = min_words - word_count
        
        # 创建统计信息工作表
        stats_sheet = wb.create_sheet(title="统计信息")
        
        # 基本统计信息
        stats_data = [
            ["总文件数", len(word_counts)],
            ["平均字数", round(sum(x[1] for x in word_counts) / len(word_counts) if word_counts else 0, 1)],
            ["最大字数", max(x[1] for x in word_counts) if word_counts else 0],
            ["最小字数", min(x[1] for x in word_counts) if word_counts else 0],
            ["字数标准", min_words],
            ["字数不足文件数", len(low_wordcount_files)],
            ["字数合格文件数", len(word_counts) - len(low_wordcount_files)],
            ["合格率", f"{(1 - len(low_wordcount_files) / len(word_counts)) * 100:.1f}%" if word_counts else "0%"]
        ]
        
        for row, (label, value) in enumerate(stats_data, 1):
            stats_sheet.cell(row=row, column=1).value = label
            stats_sheet.cell(row=row, column=2).value = value
            stats_sheet.cell(row=row, column=1).font = Font(bold=True)
        
        # 自动调整列宽
        for sheet in wb.worksheets:
            for col in range(1, sheet.max_column + 1):
                max_length = 0
                column = get_column_letter(col)
                
                for cell in sheet[column]:
                    try:
                        if cell.value and len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = (max_length + 2) * 1.2
                sheet.column_dimensions[column].width = adjusted_width
        
        # 保存工作簿
        wb.save(excel_path)
        
        return excel_path

# 添加提取作者数字的函数
def extract_author_number(filename):
    """从文件名中提取作者数字"""
    match = re.search(r'(\d+)', filename)
    return int(match.group(1)) if match else float('inf')

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()