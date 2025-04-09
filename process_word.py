from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
import sys
from zipfile import BadZipFile
from docx.oxml.ns import qn
import shutil
from docx2python import docx2python
try:
    from PIL import Image
except ImportError:
    # PIL 模块可选，用于图片处理
    pass

# 导入样式模块
from document_styles import (
    create_document_with_styles, 
    apply_title_style, 
    apply_subtitle_style, 
    apply_author_style, 
    apply_body_style,
    # 导入中文格式样式函数
    create_chinese_formal_document,
    apply_chinese_main_title,
    apply_chinese_subtitle,
    apply_chinese_body,
    apply_chinese_heading1,
    apply_chinese_heading2,
    apply_chinese_heading34
)

# 导入简单的标题识别功能
from heading_utils import is_heading1

def extract_author_number(filename):
    """从文件名中提取作者数字"""
    match = re.search(r'(\d+)', filename)
    return int(match.group(1)) if match else float('inf')

def extract_author_from_filename(filename):
    """从文件名中提取作者名"""
    try:
        # 匹配文件名中852后面的数字，然后后面的2-4个汉字
        author_match = re.search(r'852\d+[^一-龥]*([一-龥]{2,4})', filename)
        if author_match:
            return author_match.group(1).strip()
    except Exception:
        pass
    return None

def sanitize_filename(filename):
    """清理文件名中的非法字符并限制长度"""
    # Windows非法字符: \ / : * ? " < > | 以及换行符
    invalid_chars = r'\/:*?"<>|\n'
    for char in invalid_chars:
        filename = filename.replace(char, "_")
    
    # 移除重复的连续下划线
    while '__' in filename:
        filename = filename.replace('__', '_')
    
    # 移除重复的——符号
    while '————' in filename:
        filename = filename.replace('————', '——')
    while '——' in filename and '——' != filename[:2]:
        filename = filename.replace('——', '—')
    
    # 限制文件名长度为180个字符(Windows路径最大长度为260,预留一些空间给路径)
    if len(filename) > 180:
        base, ext = os.path.splitext(filename)
        filename = base[:176] + ext  # 截断名称但保留扩展名
    
    return filename.strip()

# 定义XML命名空间，避免在xpath调用中使用命名空间参数
namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}

def extract_images_from_doc(input_file, output_dir):
    """
    从Word文档中提取所有图片并保存到指定目录
    
    参数:
        input_file: Word文档路径
        output_dir: 图片输出目录
        
    返回:
        list: 提取的图片文件路径列表
    """
    try:
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        extracted_images = []
        
        # 针对.doc文件使用pywin32处理
        if input_file.lower().endswith('.doc'):
            try:
                import win32com.client
                import pythoncom
                
                # 初始化COM
                pythoncom.CoInitialize()
                
                # 创建临时目录
                temp_dir = os.path.join(output_dir, "temp_conversion")
                os.makedirs(temp_dir, exist_ok=True)
                
                # 临时docx文件路径
                temp_docx = os.path.join(temp_dir, "temp_file.docx")
                
                # 使用Word转换doc为docx
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                
                try:
                    doc = word_app.Documents.Open(os.path.abspath(input_file))
                    doc.SaveAs2(os.path.abspath(temp_docx), FileFormat=16)  # 16表示docx格式
                    doc.Close()
                    
                    # 处理转换后的docx文件
                    with open(temp_docx, 'rb') as f:
                        # 使用docx2python提取图片
                        from docx2python import docx2python
                        docx_content = docx2python(temp_docx)
                        
                        # 提取所有图片
                        image_index = 0
                        for image_data in docx_content.images:
                            if image_data[1]:  # 确保有图片数据
                                image_ext = '.png'  # 默认扩展名
                                image_path = os.path.join(output_dir, f"image_{image_index}{image_ext}")
                                with open(image_path, 'wb') as img_file:
                                    img_file.write(image_data[1])
                                extracted_images.append(image_path)
                                image_index += 1
                finally:
                    word_app.Quit()
                    pythoncom.CoUninitialize()
                    
                    # 清理临时文件
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
            except ImportError:
                print("× 错误: 处理.doc文件需要安装pywin32和docx2python")
                return []
        else:
            # 处理.docx文件
            try:
                doc = Document(input_file)
            except BadZipFile:
                print(f"× 错误：文件 '{input_file}' 可能已损坏或不是有效的Word文档")
                return []
            
            # 直接从文档关系中提取图片
            image_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        
                        # 确定文件扩展名
                        content_type = image_part.content_type
                        extension = '.jpg'  # 默认扩展名
                        if 'png' in content_type:
                            extension = '.png'
                        elif 'gif' in content_type:
                            extension = '.gif'
                        elif 'tiff' in content_type:
                            extension = '.tiff'
                        elif 'bmp' in content_type:
                            extension = '.bmp'
                        
                        # 保存图片到输出目录
                        image_filename = f"img_{image_index}{extension}"
                        image_path = os.path.join(output_dir, image_filename)
                        with open(image_path, 'wb') as f:
                            f.write(image_part.blob)
                        
                        extracted_images.append(image_path)
                        image_index += 1
                    except Exception as e:
                        print(f"× 提取图片时出错: {str(e)}")
        
        print(f"✓ 成功从文档中提取了 {len(extracted_images)} 张图片")
        return extracted_images
    except Exception as e:
        print(f"× 提取图片时出现错误: {str(e)}")
        return []

def process_doc_file(input_file, output_dir, suffix_enabled=True, 
                    suffix_text="——福州大学先进制造学院与海洋学院关工委2023年'中华魂'（毛泽东伟大精神品格）主题教育征文", 
                    use_chinese_format=False, keep_image_position=True, show_author_info=True):
    """
    处理 doc 格式的 Word 文件
    """
    try:
        # 检查依赖
        try:
            import win32com.client
            import pythoncom
        except ImportError:
            print(f"× 错误：处理doc文件需要pywin32库，请安装：pip install pywin32")
            return False
            
        # 在新线程中调用COM对象前需初始化COM
        pythoncom.CoInitialize()
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 提取作者信息
        filename = os.path.basename(input_file)
        author_num = extract_author_number(filename)
        author_name = extract_author_from_filename(filename)
        
        print(f"DEBUG: 开始处理doc文件 {input_file}")
        
        word = None
        doc = None
        
        try:
            # 创建临时目录
            temp_dir = os.path.join(output_dir, "temp")
            os.makedirs(temp_dir, exist_ok=True)
            
            # 创建 Word 应用程序 COM 对象
            word = win32com.client.Dispatch("Word.Application")
            
            # 尝试设置可见性
            try:
                word.Visible = 0
            except:
                print("! 无法设置Word应用程序可见性，继续处理...")
                pass
            
            # 打开文档
            try:
                doc = word.Documents.Open(os.path.abspath(input_file))
            except Exception as e:
                print(f"× 打开doc文件失败: {str(e)}")
                if word:
                    word.Quit()
                pythoncom.CoUninitialize()
                return False
            
            # 构建输出文件路径
            output_filename = f"temp_{author_num}.docx"
            output_file_path = os.path.join(temp_dir, output_filename)
            
            # 另存为 docx 格式
            try:
                doc.SaveAs2(os.path.abspath(output_file_path), FileFormat=16)  # 16 代表 docx 格式
            except Exception as e:
                print(f"× 转换doc文件失败: {str(e)}")
                if doc:
                    doc.Close(SaveChanges=0)
                if word:
                    word.Quit()
                pythoncom.CoUninitialize()
                return False
            
            # 关闭文档和Word应用程序
            doc.Close(SaveChanges=0)  # 不保存更改，因为已经另存为
            doc = None
            word.Quit()
            word = None
            
            print(f"✓ 已成功将 doc 文件转换为 docx: {output_filename}")
            
            # 继续处理转换后的 docx 文件
            success = process_word_file(output_file_path, output_dir, suffix_enabled, suffix_text, 
                                       use_chinese_format, keep_image_position, show_author_info)
            
            # 删除临时文件
            try:
                if os.path.exists(output_file_path):
                    os.remove(output_file_path)
            except:
                pass  # 忽略临时文件删除错误
            
            pythoncom.CoUninitialize()  # 释放COM资源
            return success
            
        except Exception as e:
            # 处理可能的错误
            print(f"× 处理doc文件时出错：{str(e)}")
            
            # 确保资源被释放
            if doc:
                try:
                    doc.Close(SaveChanges=0)
                except:
                    pass
            if word:
                try:
                    word.Quit()
                except:
                    pass
                    
            pythoncom.CoUninitialize()  # 释放COM资源
            return False
    except Exception as e:
        print(f"× 处理doc文件时出现错误：{str(e)}")
        return False

def process_word_file(input_file, output_dir, suffix_enabled=True, 
                     suffix_text="——福州大学先进制造学院与海洋学院关工委2023年'中华魂'（毛泽东伟大精神品格）主题教育征文", 
                     use_chinese_format=False, keep_image_position=True, show_author_info=True):
    """处理单个Word文件"""
    print(f"DEBUG: 开始处理文件 {input_file}")
    try:
        # 检查文件是否存在
        if not os.path.exists(input_file):
            print(f"× 错误：输入文件 '{input_file}' 不存在")
            return False

        # 创建成功文件文件夹
        success_dir = os.path.join(output_dir, "成功文件")
        no_image_dir = os.path.join(output_dir, "无图片成功文件")
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(no_image_dir, exist_ok=True)

        # 创建临时目录用于图片处理
        temp_dir = os.path.join(output_dir, "temp_images")
        os.makedirs(temp_dir, exist_ok=True)

        try:
            # 打开源文档
            source_doc = Document(input_file)
            
            # 创建目标文档
            if use_chinese_format:
                new_doc = create_chinese_formal_document()
            else:
                new_doc = create_document_with_styles()
                
            # 提取图片信息
            image_relations = {}
            for rel in source_doc.part.rels.values():
                if "image" in rel.reltype:
                    image_relations[rel.rId] = rel

            # 检查是否有图片
            has_images = bool(image_relations)
            print(f"DEBUG: 文档中包含 {len(image_relations)} 张图片")

            # 提取信息变量
            author_name = ""
            original_title = ""
            title_found = False
            author_added = False
            used_default_author = False
            
            # 从文件名中提取作者名
            filename = os.path.basename(input_file)
            author_name = extract_author_from_filename(filename)
            
            # 保存每个段落对应的图片，用于后续处理
            paragraph_images = {}
            current_para_index = 0
            
            # 首先遍历文档，找出所有带图片的段落
            if keep_image_position and has_images:
                print("DEBUG: 分析文档结构以保持图片位置...")
                
                # 创建临时图片目录
                images_temp_dir = os.path.join(temp_dir, "images_temp")
                os.makedirs(images_temp_dir, exist_ok=True)
                
                # 遍历所有段落，记录图片位置
                for para_idx, para in enumerate(source_doc.paragraphs):
                    # 检查段落是否包含图片
                    para_images = []
                    
                    # 遍历所有run
                    for run in para.runs:
                        # 检查run是否包含图片
                        try:
                            run_xml = run._element.xml
                            if '<a:blip' in run_xml:  # 简单检查是否有图片元素
                                # 提取所有图片关系ID
                                import re
                                rel_matches = re.findall(r'r:embed="(rId\d+)"', run_xml)
                                
                                for rel_id in rel_matches:
                                    if rel_id in image_relations:
                                        # 保存图片到临时文件
                                        try:
                                            image_part = image_relations[rel_id].target_part
                                            img_temp_path = os.path.join(images_temp_dir, f"{rel_id}.png")
                                            
                                            with open(img_temp_path, 'wb') as img_file:
                                                img_file.write(image_part.blob)
                                                
                                            para_images.append(img_temp_path)
                                            print(f"DEBUG: 在段落 {para_idx} 中找到图片 {rel_id}")
                                        except Exception as e:
                                            print(f"× 处理图片 {rel_id} 时出错: {str(e)}")
                        except Exception as e:
                            print(f"× 检查run中的图片时出错: {str(e)}")
                    
                    # 如果段落包含图片，记录下来
                    if para_images:
                        paragraph_images[para_idx] = para_images
            
            # 处理文档内容
            for para_idx, para in enumerate(source_doc.paragraphs):
                try:
                    text = para.text.strip()
                    
                    # 对于空段落，检查是否有图片
                    if not text and para_idx in paragraph_images:
                        # 创建新段落并插入图片
                        img_para = new_doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                        for img_path in paragraph_images[para_idx]:
                            try:
                                run = img_para.add_run()
                                run.add_picture(img_path, width=Inches(6))
                                print(f"DEBUG: 在空段落 {para_idx} 中添加图片")
                            except Exception as e:
                                print(f"× 添加图片时出错: {str(e)}")
                        
                        continue  # 处理完图片，继续下一段落
                    
                    if not text:
                        continue  # 跳过空段落
                    
                    # 提取标题（第一个非空段落）
                    if not title_found:
                        original_title = text
                        title_para = new_doc.add_paragraph(original_title)
                        if use_chinese_format:
                            apply_chinese_main_title(new_doc, title_para)
                        else:
                            apply_title_style(new_doc, title_para)
                        
                        # 添加副标题
                        if suffix_enabled and suffix_text:
                            subtitle_para = new_doc.add_paragraph(suffix_text)
                            if use_chinese_format:
                                apply_chinese_subtitle(new_doc, subtitle_para)
                            else:
                                apply_subtitle_style(new_doc, subtitle_para)
                        
                        title_found = True
                        continue
                    
                    # 从文档内容中提取作者信息
                    if not author_name and text.startswith('852'):
                        author_match = re.search(r'852\d*[^-]*-([^-\d\W]+)', text)
                        if not author_match:
                            author_match = re.search(r'852\d*[\s-]*([^\d\W]+)', text)
                        if author_match:
                            author_name = author_match.group(1).strip()
                        continue  # 跳过学号行
                    
                    # 添加作者信息（仅在需要且未添加时）
                    if author_name and not author_added and not text.startswith('852') and show_author_info:
                        author_text = f"（先进制造学院与海洋学院关工委通讯员{author_name}）"
                        author_para = new_doc.add_paragraph(author_text)
                        if use_chinese_format:
                            apply_chinese_subtitle(new_doc, author_para)
                        else:
                            apply_author_style(new_doc, author_para)
                        author_added = True
                        continue
                    
                    # 处理正文段落
                    if title_found and not text.startswith('852'):
                        # 添加文本段落
                        body_para = new_doc.add_paragraph(text)
                        if use_chinese_format:
                            apply_chinese_body(new_doc, body_para)
                        else:
                            apply_body_style(new_doc, body_para)
                        
                        # 如果段落有关联图片且需要保持图片位置，则添加图片
                        if keep_image_position and para_idx in paragraph_images:
                            # 添加段落后面的图片
                            img_para = new_doc.add_paragraph()
                            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            for img_path in paragraph_images[para_idx]:
                                try:
                                    run = img_para.add_run()
                                    run.add_picture(img_path, width=Inches(6))
                                    print(f"DEBUG: 在段落 {para_idx} 后添加图片")
                                except Exception as e:
                                    print(f"× 添加图片时出错: {str(e)}")
                except Exception as e:
                    print(f"× 处理段落 {para_idx} 时出错：{str(e)}")
                    continue

            # 处理默认作者名
            if not author_name:
                author_name = "佚名"
                used_default_author = True
                print(f"! 警告：未能提取作者名，使用默认值\"{author_name}\"")
                
                # 添加默认作者信息（如果需要且未添加）
                if not author_added and show_author_info and title_found:
                    author_text = f"（先进制造学院与海洋学院关工委通讯员{author_name}）"
                    author_para = new_doc.add_paragraph(author_text)
                    if use_chinese_format:
                        apply_chinese_subtitle(new_doc, author_para)
                    else:
                        apply_author_style(new_doc, author_para)
                    author_added = True

            # 如果不保持图片位置或没找到图片位置信息，但文档包含图片，则在末尾添加图片
            if has_images and not keep_image_position:
                print("DEBUG: 将所有图片添加到文档末尾")
                new_doc.add_paragraph()  # 添加空行分隔
                
                # 将所有图片提取并添加到末尾
                temp_img_dir = os.path.join(temp_dir, "all_images")
                os.makedirs(temp_img_dir, exist_ok=True)
                
                image_files = extract_images_from_doc(input_file, temp_img_dir)
                
                for img_path in image_files:
                    try:
                        img_para = new_doc.add_paragraph()
                        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = img_para.add_run()
                        run.add_picture(img_path, width=Inches(6))
                    except Exception as e:
                        print(f"× 添加图片到文档末尾时出错：{str(e)}")

            # 生成输出文件名
            if original_title:
                # 防止副标题在原标题中重复出现
                if suffix_enabled and suffix_text and suffix_text in original_title:
                    original_title = original_title.split(suffix_text)[0].strip()
                
                # 根据用户配置生成文件名
                if show_author_info:
                    if suffix_enabled and suffix_text:
                        new_filename = f"({author_name}){original_title}{suffix_text}.docx"
                    else:
                        new_filename = f"({author_name}){original_title}.docx"
                else:
                    # 不添加作者信息时，文件名中也不包含作者名
                    if suffix_enabled and suffix_text:
                        new_filename = f"{original_title}{suffix_text}.docx"
                    else:
                        new_filename = f"{original_title}.docx"
                
                # 清理文件名
                new_filename = sanitize_filename(new_filename)
                # 根据是否有图片选择保存目录
                output_dir_final = success_dir if has_images else no_image_dir
                output_file = os.path.join(output_dir_final, new_filename)
                
                # 保存文档
                try:
                    new_doc.save(output_file)
                    if used_default_author:
                        if has_images:
                            print(f"✓ 文件处理完成（使用默认作者名）：{new_filename}")
                        else:
                            print(f"✓ 文件处理完成（使用默认作者名，无图片）：{new_filename}")
                    else:
                        if has_images:
                            print(f"✓ 文件处理完成：{new_filename}")
                        else:
                            print(f"✓ 文件处理完成（无图片）：{new_filename}")
                except Exception as e:
                    print(f"× 保存文件时出错：{str(e)}")
                    return False
            else:
                print("× 未能提取标题")
                return False
                
        except BadZipFile:
            print(f"× 错误：文件 '{input_file}' 可能已损坏或不是有效的Word文档")
            return False

        # 清理临时文件
        try:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"! 清理临时文件时出错: {str(e)}")

        return True

    except Exception as e:
        print(f"× 处理文件时出现错误：{str(e)}")
        return False

def process_folder(input_folder, output_folder):
    """处理文件夹中的所有Word文档"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    temp_dir = os.path.join(output_folder, "temp_images")
    
    try:
        for filename in os.listdir(input_folder):
            if filename.startswith('~$') or not filename.endswith('.docx'):
                continue
                
            if hasattr(sys.stdout, 'set_current_file'):
                sys.stdout.set_current_file(filename)
            if hasattr(sys.stderr, 'set_current_file'):
                sys.stderr.set_current_file(filename)
                
            input_path = os.path.join(input_folder, filename)
            print(f"\n处理文件：{filename}")
            
            # 处理文档
            process_word_file(input_path, output_folder)
                    
    except Exception as e:
        print(f"× 处理文件夹时发生错误：{str(e)}")
    finally:
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
        if hasattr(sys.stdout, 'set_current_file'):
            sys.stdout.set_current_file(None)
        if hasattr(sys.stderr, 'set_current_file'):
            sys.stderr.set_current_file(None)

# 使用示例
if __name__ == "__main__":
    input_folder = r"C:\Users\86159\Desktop\新建文件夹 (2)"
    output_folder = r"C:\Users\86159\Desktop\新建文件夹 (2)"
    
    process_folder(input_folder, output_folder)