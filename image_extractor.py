"""
提供从Word文档中提取图片的功能
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
import sys
import shutil
from zipfile import BadZipFile
try:
    from PIL import Image
except ImportError:
    # PIL 模块可选，用于图片处理
    pass

def extract_images_from_doc(input_file, output_dir):
    """
    从Word文档中提取所有图片并保存到指定目录
    
    参数:
        input_file: Word文档路径
        output_dir: 图片输出目录
        
    返回:
        list: 提取的图片文件路径列表
    """
    try:
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        extracted_images = []
        
        # 针对.doc文件使用pywin32处理
        if input_file.lower().endswith('.doc'):
            try:
                import win32com.client
                import pythoncom
                
                # 初始化COM
                pythoncom.CoInitialize()
                
                # 创建临时目录
                temp_dir = os.path.join(output_dir, "temp_conversion")
                os.makedirs(temp_dir, exist_ok=True)
                
                # 临时docx文件路径
                temp_docx = os.path.join(temp_dir, "temp_file.docx")
                
                # 使用Word转换doc为docx
                word_app = win32com.client.Dispatch("Word.Application")
                word_app.Visible = False
                
                try:
                    doc = word_app.Documents.Open(os.path.abspath(input_file))
                    doc.SaveAs2(os.path.abspath(temp_docx), FileFormat=16)  # 16表示docx格式
                    doc.Close()
                    
                    # 处理转换后的docx文件
                    with open(temp_docx, 'rb') as f:
                        # 使用docx2python提取图片
                        from docx2python import docx2python
                        docx_content = docx2python(temp_docx)
                        
                        # 提取所有图片
                        image_index = 0
                        for image_data in docx_content.images:
                            if image_data[1]:  # 确保有图片数据
                                image_ext = '.png'  # 默认扩展名
                                image_path = os.path.join(output_dir, f"image_{image_index}{image_ext}")
                                with open(image_path, 'wb') as img_file:
                                    img_file.write(image_data[1])
                                extracted_images.append(image_path)
                                image_index += 1
                finally:
                    word_app.Quit()
                    pythoncom.CoUninitialize()
                    
                    # 清理临时文件
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
            except ImportError:
                print("× 错误: 处理.doc文件需要安装pywin32和docx2python")
                return []
        else:
            # 处理.docx文件
            try:
                doc = Document(input_file)
            except BadZipFile:
                print(f"× 错误：文件 '{input_file}' 可能已损坏或不是有效的Word文档")
                return []
            
            # 直接从文档关系中提取图片
            image_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        
                        # 确定文件扩展名
                        content_type = image_part.content_type
                        extension = '.jpg'  # 默认扩展名
                        if 'png' in content_type:
                            extension = '.png'
                        elif 'gif' in content_type:
                            extension = '.gif'
                        elif 'tiff' in content_type:
                            extension = '.tiff'
                        elif 'bmp' in content_type:
                            extension = '.bmp'
                        
                        # 保存图片到输出目录
                        image_filename = f"img_{image_index}{extension}"
                        image_path = os.path.join(output_dir, image_filename)
                        with open(image_path, 'wb') as f:
                            f.write(image_part.blob)
                        
                        extracted_images.append(image_path)
                        image_index += 1
                    except Exception as e:
                        print(f"× 提取图片时出错: {str(e)}")
        
        print(f"✓ 成功从文档中提取了 {len(extracted_images)} 张图片")
        return extracted_images
    except Exception as e:
        print(f"× 提取图片时出现错误: {str(e)}")
        return []

def add_images_to_document(doc, images, width=None):
    """
    将图片添加到文档末尾
    
    参数:
        doc: Document对象
        images: 图片路径列表
        width: 图片宽度，默认为6英寸
    """
    if not images:
        return
        
    doc.add_paragraph()  # 添加空行分隔
    
    width = width or Inches(6)  # 默认宽度为6英寸
    
    for img_path in images:
        try:
            img_para = doc.add_paragraph()
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = img_para.add_run()
            run.add_picture(img_path, width=width)
        except Exception as e:
            print(f"× 添加图片到文档时出错：{str(e)}")
            
def extract_document_image_relations(doc):
    """
    提取文档中的图片关系
    
    参数:
        doc: Document对象
        
    返回:
        dict: 图片关系ID到关系对象的映射
    """
    image_relations = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_relations[rel.rId] = rel
    
    return image_relations

def find_paragraph_images(doc, image_relations):
    """
    找出文档中每个段落关联的图片
    
    参数:
        doc: Document对象
        image_relations: 图片关系字典
        
    返回:
        dict: 段落索引到图片路径列表的映射
    """
    paragraph_images = {}
    
    # 创建临时图片目录
    temp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_images")
    os.makedirs(temp_dir, exist_ok=True)
    
    # 遍历所有段落
    for para_idx, para in enumerate(doc.paragraphs):
        # 检查段落是否包含图片
        para_images = []
        
        # 遍历所有run
        for run in para.runs:
            # 检查run是否包含图片
            try:
                run_xml = run._element.xml
                if '<a:blip' in run_xml:  # 简单检查是否有图片元素
                    # 提取所有图片关系ID
                    import re
                    rel_matches = re.findall(r'r:embed="(rId\d+)"', run_xml)
                    
                    for rel_id in rel_matches:
                        if rel_id in image_relations:
                            # 保存图片到临时文件
                            try:
                                image_part = image_relations[rel_id].target_part
                                img_temp_path = os.path.join(temp_dir, f"{rel_id}.png")
                                
                                with open(img_temp_path, 'wb') as img_file:
                                    img_file.write(image_part.blob)
                                    
                                para_images.append(img_temp_path)
                            except Exception as e:
                                print(f"× 处理图片 {rel_id} 时出错: {str(e)}")
            except Exception as e:
                print(f"× 检查run中的图片时出错: {str(e)}")
        
        # 如果段落包含图片，记录下来
        if para_images:
            paragraph_images[para_idx] = para_images
    
    return paragraph_images